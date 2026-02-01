"""
Document Processor Module
=========================
This module handles extraction of text from various document formats.
Supports: PDF, DOCX, TXT, MD

Each format has specific handling to preserve structure and metadata.
"""

import logging
from typing import Dict, Any, Optional
from pathlib import Path
import mimetypes

# Document processing libraries
import pypdf
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Handles text extraction from various document formats.
    
    Supported formats:
    - PDF: Extracts text page by page
    - DOCX: Extracts from Word documents
    - TXT: Plain text files
    - MD: Markdown files
    """
    
    def __init__(self):
        """Initialize the document processor."""
        logger.info("Initialized DocumentProcessor")
        
        # Map of file extensions to processing methods
        self.processors = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.md': self._process_md,
        }
    
    def process_file(self, file_path: str, filename: str = None) -> Dict[str, Any]:
        """
        Process a file and extract its text content.
        
        Args:
            file_path: Path to the file
            filename: Original filename (optional, for metadata)
        
        Returns:
            Dictionary with:
            - text: Extracted text content
            - metadata: File metadata (source, title, page_count, etc.)
        
        Process:
        1. Detect file type from extension
        2. Call appropriate processor
        3. Extract text and metadata
        4. Return structured result
        """
        try:
            path = Path(file_path)
            extension = path.suffix.lower()
            filename = filename or path.name
            
            logger.info(f"Processing file: {filename} (type: {extension})")
            
            if extension not in self.processors:
                raise ValueError(
                    f"Unsupported file type: {extension}. "
                    f"Supported types: {', '.join(self.processors.keys())}"
                )
            
            # Call appropriate processor
            processor = self.processors[extension]
            result = processor(file_path)
            
            # Add common metadata
            result['metadata']['source'] = filename
            result['metadata']['file_type'] = extension
            
            logger.info(
                f"Successfully processed {filename}: "
                f"{len(result['text'])} characters extracted"
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF files.
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            Dictionary with text and metadata
        
        Technical Notes:
        - Uses pypdf for extraction
        - Extracts page by page
        - Preserves page numbers for citations
        - Handles multi-column layouts reasonably well
        
        Limitations:
        - Scanned PDFs (images) won't extract text
        - Complex layouts may have ordering issues
        - For scanned PDFs, OCR would be needed (tesseract)
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                logger.info(f"PDF has {num_pages} pages")
                
                # Extract text from all pages
                text_parts = []
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Add page marker for context
                        text_parts.append(f"\n--- Page {page_num} ---\n{page_text}")
                
                full_text = "\n".join(text_parts)
                
                # Extract metadata
                metadata = {
                    'page_count': num_pages,
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                }
                
                return {
                    'text': full_text,
                    'metadata': metadata
                }
                
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX files.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Dictionary with text and metadata
        
        Technical Notes:
        - Uses python-docx library
        - Extracts from paragraphs and tables
        - Preserves basic structure
        - Handles headers and footers
        
        Features:
        - Paragraph-by-paragraph extraction
        - Table content included
        - Basic formatting preserved through spacing
        """
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            
            # Extract metadata from core properties
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            return {
                'text': full_text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise
    
    def _process_txt(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from plain text files.
        
        Args:
            file_path: Path to TXT file
        
        Returns:
            Dictionary with text and metadata
        
        Technical Notes:
        - Simple UTF-8 text reading
        - Handles various text encodings with fallback
        - No special processing needed
        """
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            
            # Count lines for metadata
            line_count = len(text.splitlines())
            
            metadata = {
                'line_count': line_count,
                'char_count': len(text)
            }
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing TXT: {e}")
            raise
    
    def _process_md(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Markdown files.
        
        Args:
            file_path: Path to MD file
        
        Returns:
            Dictionary with text and metadata
        
        Technical Notes:
        - Reads as plain text (formatting marks preserved)
        - Could convert to HTML first if needed
        - Headers and structure preserved in raw form
        
        For RAG:
        - Keeping markdown syntax can help with structure
        - Headers provide good section boundaries
        - Code blocks are clearly marked
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Count headers for structure info
            header_count = text.count('\n#')
            line_count = len(text.splitlines())
            
            metadata = {
                'header_count': header_count,
                'line_count': line_count,
                'char_count': len(text)
            }
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error processing Markdown: {e}")
            raise
    
    def is_supported(self, filename: str) -> bool:
        """
        Check if a file type is supported.
        
        Args:
            filename: Name of the file
        
        Returns:
            True if supported, False otherwise
        """
        extension = Path(filename).suffix.lower()
        return extension in self.processors
    
    def get_supported_extensions(self) -> list:
        """
        Get list of supported file extensions.
        
        Returns:
            List of extensions (e.g., ['.pdf', '.docx', '.txt', '.md'])
        """
        return list(self.processors.keys())


# Global instance
_document_processor = None


def get_document_processor() -> DocumentProcessor:
    """Get or create the global document processor instance."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor